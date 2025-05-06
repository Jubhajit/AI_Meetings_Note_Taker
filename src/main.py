import gradio as gr
import pyaudio
import wave
import threading
import whisper
import noisereduce as nr
import numpy as np
from transformers import pipeline
import tempfile
import torch
import os
from langdetect import detect
import soundfile as sf
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Audio Recording Variables
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 44100
CHUNK = 1024
recording = False
frames = []
transcription_result = ""
summary_result = ""

def start_recording():
    global recording, frames
    recording = True
    frames = []

    def record():
        audio = pyaudio.PyAudio()
        stream = audio.open(format=FORMAT, channels=CHANNELS, rate=RATE, input=True, frames_per_buffer=CHUNK)
        while recording:
            data = stream.read(CHUNK)
            frames.append(data)
        stream.stop_stream()
        stream.close()
        audio.terminate()
        process_audio()

    threading.Thread(target=record, daemon=True).start()
    return "Recording in progress...", gr.update(visible=True), gr.update(visible=False), gr.update(variant="primary")

def stop_recording():
    global recording
    recording = False
    return "Processing... Please wait.", gr.update(visible=False), gr.update(visible=True), gr.update(variant="secondary")

def process_audio():
    global transcription_result, summary_result, frames

    transcription_result = ""
    summary_result = ""

    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"Using device: {device}")

    audio = pyaudio.PyAudio()

    with tempfile.NamedTemporaryFile(suffix=".flac", delete=False) as tmpfile:
        audio_data = np.frombuffer(b"".join(frames), dtype=np.int16)
        reduced_noise = nr.reduce_noise(y=audio_data, sr=RATE)
        sf.write(tmpfile.name, reduced_noise, RATE, format='FLAC')

    try:
        model = whisper.load_model("small").to(device)
        transcription_result = transcribe_audio(tmpfile.name, model)
        summary_result = summarize_text(transcription_result)

    except RuntimeError as e:
        if "CUDA error: device-side assert triggered" in str(e):
            print("CUDA error detected. Switching to CPU...")

            torch.cuda.empty_cache()
            torch.cuda.ipc_collect()

            device = "cpu"
            model = whisper.load_model("small").to(device)

            transcription_result = transcribe_audio(tmpfile.name, model)
            summary_result = summarize_text(transcription_result)

    return transcription_result, summary_result

def transcribe_audio(audio_file, model):
    result = model.transcribe(audio_file, language=None)
    transcription_with_timestamps = "\n".join([
        f"[{segment['start']:.2f}s - {segment['end']:.2f}s] {segment['text']}"
        for segment in result["segments"]])

    return transcription_with_timestamps

def remove_duplicates(summary_text):
    lines = summary_text.split("\n")
    return "\n".join(list(dict.fromkeys(lines)))

def summarize_text(transcript):
    cleaned_transcript = re.sub(r"\[\d+\.\d+s - \d+\.\d+s\]", "", transcript).strip()
    sentences = cleaned_transcript.split(". ")
    reduced_text = ". ".join(sentences[:len(sentences) // 3])

    if len(reduced_text) < 200:
        reduced_text = ". ".join(sentences[:len(sentences)])

    summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=-1)

    try:
        summary = summarizer(
            reduced_text,
            max_length=400,
            min_length=100,
            do_sample=False,
            num_beams=8,
            truncation=True
        )[0]["summary_text"]
    except Exception as e:
        print(f"Error during summarization: {e}")
        summary = "Summary unavailable due to an error."

    summary = remove_duplicates(summary)
    bullet_points = "\n".join([f"- {point.strip()}" for point in summary.split(". ") if len(point) > 5])

    return bullet_points

def generate_docx():
    global transcription_result, summary_result
    docx_filename = "meeting_notes.docx"
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'DejaVu Sans'
    font.size = Pt(11)

    document.add_heading("📌 Meeting Notes", level=1)

    document.add_paragraph("📝 Meeting Title: [Add Title]")
    document.add_paragraph("📍 बैठक स्थान (Meeting Location): [स्थान जोड़ें]")
    document.add_paragraph("👥 उपस्थित लोग (Attendees): [उपस्थित लोगों को जोड़ें]")

    document.add_paragraph("\n")

    document.add_heading("🔊 Transcription:", level=2)
    lines_per_page = 40
    line_count = 0
    for line in transcription_result.split('\n'):
        if line.strip():
            document.add_paragraph(line.strip())
            line_count += 1
            if line_count >= lines_per_page:
                document.add_page_break()
                line_count = 0

    document.add_page_break()
    document.add_heading("📄 Summary:", level=2)
    line_count = 0
    for line in summary_result.split('\n'):
        if line.strip():
            document.add_paragraph(line.strip())
            line_count += 1
            if line_count >= lines_per_page:
                document.add_page_break()
                line_count = 0

    document.save(docx_filename)
    return docx_filename

def download_docx():
    global transcription_result, summary_result
    docx_path = generate_docx()
    return docx_path

with gr.Blocks(theme=gr.themes.Soft()) as app:
    gr.Markdown(
        """
        # 🎙️ AI-Meeting Minutes
        ### Automatically Transcribe and Summarize your meetings(Both Sides) with AI 
         Works with Zoom , Meet , Skype , Teams , Whatsapp call...
        """
    )

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### Recording Controls")
            with gr.Group():
                with gr.Row():
                    start_btn = gr.Button("🎙️ Start Recording", variant="primary", size="lg",visible=True)
                    stop_btn = gr.Button("⏹️ Stop Recording", variant="stop", size="lg",visible=True)

                status_display = gr.Textbox(
                    value="Ready to record your meeting",
                    label="Status",
                    interactive=False
                )

                with gr.Row():
                    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

                    record_gif_path = os.path.join(BASE_DIR, "assets", "record.gif")

                    record_indicator = gr.Image(
                        record_gif_path, label="Recording in progress",
                        visible=False,
                        show_download_button=False,
                        container=False,
                        height=100,
                        width=100)

        with gr.Column(scale=2):
            gr.Markdown("### Processing Status")
            with gr.Accordion("How to use", open=False):
                gr.Markdown("""
                1. Click "Start Recording" to begin capturing audio
                2. Records Both Side of the Meeting(with Noise Reduction)
                3. Click "Stop Recording" when finished
                4. View the transcription and summary in their respective tabs
                5. Download the complete notes as a DOCX file

                6.If using BLUETOOTH will record only MIC

                7.Optimized to Use both GPU & CPU

                8.Records 10-15 Mins at a time without ERROR
                """)

            progress = gr.Progress(track_tqdm=True)

    with gr.Tabs() as tabs:
        with gr.TabItem("✍️ Transcription", id="transcription"):
            transcript_output = gr.Textbox(
                label="Full Meeting Transcription",
                placeholder="Transcription will appear here after recording...",
                interactive=False,
                lines=15
            )

        with gr.TabItem("📝 Summary", id="summary"):
            summary_output = gr.Textbox(
                label="Meeting Summary",
                placeholder="AI-generated summary will appear here after recording...",
                interactive=False,
                lines=10
            )

    with gr.Row():
        download_btn = gr.Button("📥 Download DOCX Report", variant="secondary", size="lg")
        download_file = gr.File(label="Download Complete Meeting Notes", visible=True)

    with gr.Accordion("About ", open=False):
        gr.Markdown("""
        **MADE BY -JUBHAJIT DEB** (AI Minutes of Meeting) uses advanced AI to transcribe and summarize your meetings.

        **Features:**
        - Records Meeting on Both Side(Skype, Zoom , Meets , Teams , Whatsapp call)
        - Accurate speech-to-text transcription using Whisper AI
        - Intelligent summarization with BART
        - Multi-language support(en & hi)
        - Complete DOCX report generation

        This tool helps you focus on your meeting while AI takes notes for you!
        """)

    start_btn.click(
        start_recording,
        [],
        [status_display, record_indicator, start_btn, stop_btn]
    )

    stop_btn.click(
        stop_recording,
        [],
        [status_display, record_indicator, stop_btn, start_btn]
    )

    stop_btn.click(
        process_audio,
        [],
        [transcript_output, summary_output]
    )

    download_btn.click(
        download_docx,
        [],
        [download_file]
    )

if __name__ == "__main__":
    app.queue().launch()
